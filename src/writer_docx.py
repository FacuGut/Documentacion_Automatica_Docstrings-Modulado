"""
Módulo de escritura de documentos `.docx`.
- Generar el documento docx a partir de los resultados de core (lista de DocItem).
- Aplicar estilos y colores en títulos según si hay docstring o no.
- No decide qué documentar: solo renderiza lo que le pasa app/orquestador.py.
"""
import os
from typing import Dict, List
from docx import Document
from docx.shared import RGBColor
from core.modelos import DocItem

def escribir_docx_por_archivo(
    resultados: Dict[str, List[DocItem]],
    ruta_raiz: str,
    salida_docx: str,
    color_con_doc: tuple[int, int, int],
    color_sin_doc: tuple[int, int, int],
) -> None:
    doc = Document()
    doc.add_heading("Documentación de funciones con docstrings", level=0)

    azul = RGBColor(*color_con_doc)
    rojo = RGBColor(*color_sin_doc)

    for path_abs, items in resultados.items():
        if not items:
            continue

        relativo = os.path.relpath(path_abs, ruta_raiz)
        doc.add_heading(relativo, level=1)

        for it in items:
            heading = doc.add_heading(level=2)
            run = heading.add_run(it.nombre)
            run.font.bold = True
            run.font.color.rgb = azul if it.docstring else rojo

            if not it.docstring:
                doc.add_paragraph(
                    "La clase no tiene descripción." if it.tipo == "class"
                    else "La función no tiene descripción."
                )
            else:
                for linea in it.docstring.splitlines():
                    doc.add_paragraph(linea)

    doc.save(salida_docx)